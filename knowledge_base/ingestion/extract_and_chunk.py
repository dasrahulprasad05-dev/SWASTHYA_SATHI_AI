"""
Swasthya Sathi AI — DOCX Knowledge Base Extractor
Extracts text from all .docx files, cleans it, chunks it with metadata,
and outputs structured JSON knowledge files for RAG ingestion.
"""

import sys
sys.path.insert(0, r"D:\SWASTHYA SATHI AI\knowledge_base\ingestion\.pip_packages")

import os
import json
import re
from docx import Document

RAW_DOCS_DIRS = [
    r"D:\SWASTHYA SATHI AI\knowledge_base\raw_docs\rahul knowladge base",
    r"D:\SWASTHYA SATHI AI\knowledge_base\raw_docs\knowladge base 2",
]
OUTPUT_DIR = r"D:\SWASTHYA SATHI AI\knowledge_base\processed"
CHUNKS_OUTPUT = r"D:\SWASTHYA SATHI AI\knowledge_base\chunks"

os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(CHUNKS_OUTPUT, exist_ok=True)

# Category mapping based on filename keywords
CATEGORY_MAP = {
    'anemia': 'diseases',
    'diabetes': 'diseases',
    'hypertension': 'diseases',
    'chickenpox': 'diseases',
    'dengue': 'diseases',
    'malaria': 'diseases',
    'typhoid': 'diseases',
    'tuberculosis': 'diseases',
    'jaundice': 'diseases',
    'hepatitis': 'diseases',
    'viral_fever': 'diseases',
    'conjunctivitis': 'diseases',
    'skin_infection': 'diseases',
    'cold_symptom': 'symptoms',
    'cough_symptom': 'symptoms',
    'vomiting': 'symptoms',
    'common_cold': 'diseases',
    'flu': 'diseases',
    'pneumonia': 'child_health',
    'diarrhoea': 'child_health',
    'dehydration': 'child_health',
    'monsoon': 'prevention',
    'hygiene': 'prevention',
    'sanitation': 'prevention',
    'nutrition': 'prevention',
    'child_health': 'child_health',
    'immunization': 'child_health',
    'maternal': 'maternal_health',
    'mamata': 'government_schemes',
    'janani': 'government_schemes',
    'bsky': 'government_schemes',
    'ayushman': 'government_schemes',
    'pmjay': 'government_schemes',
    'elderly': 'elderly_health',
    'emergency': 'emergency_first_aid',
    'warning': 'emergency_first_aid',
    'khordha': 'hospitals',
    'cuttack': 'hospitals',
    'puri': 'hospitals',
    'district': 'hospitals',
    'healthcare_facilities': 'hospitals',
}

def detect_category(filename: str) -> str:
    fn_lower = filename.lower()
    for keyword, category in CATEGORY_MAP.items():
        if keyword in fn_lower:
            return category
    return 'general'

def extract_text_from_docx(filepath: str) -> str:
    """Extract all text from a .docx file."""
    doc = Document(filepath)
    full_text = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            full_text.append(text)
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                full_text.append(row_text)
    return '\n'.join(full_text)

def clean_text(text: str) -> str:
    """Clean extracted text while preserving Odia, Hindi & English characters."""
    # Remove excessive whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]{2,}', ' ', text)
    # Remove page numbers & artifacts
    text = re.sub(r'Page \d+ of \d+', '', text)
    return text.strip()

def chunk_text(text: str, chunk_size: int = 800, overlap: int = 100) -> list:
    """Split text into overlapping chunks for vector embedding."""
    sentences = re.split(r'(?<=[.!?।\n])\s+', text)
    chunks = []
    current_chunk = []
    current_length = 0

    for sentence in sentences:
        sentence_length = len(sentence)
        if current_length + sentence_length > chunk_size and current_chunk:
            chunk_text_str = ' '.join(current_chunk)
            chunks.append(chunk_text_str)
            # Keep overlap
            overlap_tokens = []
            overlap_len = 0
            for s in reversed(current_chunk):
                if overlap_len + len(s) <= overlap:
                    overlap_tokens.insert(0, s)
                    overlap_len += len(s)
                else:
                    break
            current_chunk = overlap_tokens
            current_length = overlap_len
        current_chunk.append(sentence)
        current_length += sentence_length

    if current_chunk:
        chunks.append(' '.join(current_chunk))

    return chunks

def detect_language(text: str) -> str:
    """Detect dominant language in text."""
    odia_chars = len(re.findall(r'[\u0B00-\u0B7F]', text))
    hindi_chars = len(re.findall(r'[\u0900-\u097F]', text))
    english_chars = len(re.findall(r'[a-zA-Z]', text))
    
    total = odia_chars + hindi_chars + english_chars
    if total == 0:
        return 'en'
    
    if odia_chars / total > 0.3:
        return 'or'
    elif hindi_chars / total > 0.3:
        return 'hi'
    return 'en'

def process_all_documents():
    """Main pipeline: Extract -> Clean -> Chunk -> Save."""
    all_chunks = []
    processed_files = []

    # Collect all .docx files from all source directories
    all_docx = []
    for docs_dir in RAW_DOCS_DIRS:
        if os.path.isdir(docs_dir):
            for f in os.listdir(docs_dir):
                if f.endswith('.docx'):
                    all_docx.append((docs_dir, f))
    print(f"Found {len(all_docx)} .docx files to process\n")

    for docs_dir, filename in sorted(all_docx, key=lambda x: x[1]):
        filepath = os.path.join(docs_dir, filename)
        doc_id = os.path.splitext(filename)[0]
        category = detect_category(filename)
        
        print(f"[DOC] Processing: {filename}")
        print(f"   Category: {category}")

        # 1. Extract
        raw_text = extract_text_from_docx(filepath)
        
        # 2. Clean
        cleaned_text = clean_text(raw_text)
        language = detect_language(cleaned_text)
        print(f"   Language: {language} | Characters: {len(cleaned_text)}")

        # 3. Save full cleaned document
        doc_metadata = {
            'id': doc_id,
            'filename': filename,
            'category': category,
            'language': language,
            'char_count': len(cleaned_text),
            'content': cleaned_text,
        }
        
        doc_output_path = os.path.join(OUTPUT_DIR, f"{doc_id}.json")
        with open(doc_output_path, 'w', encoding='utf-8') as f:
            json.dump(doc_metadata, f, ensure_ascii=False, indent=2)

        # 4. Chunk
        chunks = chunk_text(cleaned_text)
        print(f"   Chunks: {len(chunks)}")

        for i, chunk in enumerate(chunks):
            chunk_entry = {
                'chunk_id': f"{doc_id}_chunk_{i}",
                'document_id': doc_id,
                'document_title': doc_id.replace('_', ' '),
                'category': category,
                'language': language,
                'chunk_index': i,
                'total_chunks': len(chunks),
                'content': chunk,
                'char_count': len(chunk),
            }
            all_chunks.append(chunk_entry)

        processed_files.append({
            'filename': filename,
            'category': category,
            'language': language,
            'chunks': len(chunks),
            'characters': len(cleaned_text),
        })
        print()

    # 5. Save all chunks as a single JSON index
    chunks_index_path = os.path.join(CHUNKS_OUTPUT, 'all_chunks.json')
    with open(chunks_index_path, 'w', encoding='utf-8') as f:
        json.dump(all_chunks, f, ensure_ascii=False, indent=2)

    # 6. Save processing manifest
    manifest = {
        'total_documents': len(processed_files),
        'total_chunks': len(all_chunks),
        'documents': processed_files,
    }
    manifest_path = os.path.join(CHUNKS_OUTPUT, 'manifest.json')
    with open(manifest_path, 'w', encoding='utf-8') as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)

    print("=" * 60)
    print(f"[OK] PROCESSING COMPLETE")
    print(f"   Documents: {len(processed_files)}")
    print(f"   Total Chunks: {len(all_chunks)}")
    print(f"   Output: {OUTPUT_DIR}")
    print(f"   Chunks Index: {chunks_index_path}")
    print("=" * 60)

if __name__ == '__main__':
    process_all_documents()
